import asyncio
import hashlib
import io
import json
import mimetypes
import shutil
import uuid
from pathlib import Path

import fitz
from docx import Document
from fastapi import UploadFile
from PIL import Image
from pypdf import PdfReader

from app.config import Settings
from app.models import OcrEngineName
from app.ocr.service import OcrService

IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}
TEXT_SUFFIXES = {".txt", ".md", ".csv", ".json"}


async def save_upload(file: UploadFile, upload_dir: Path) -> Path:
    safe_name = Path(file.filename or "upload.bin").name
    target = upload_dir / f"{uuid.uuid4().hex}_{safe_name}"
    with target.open("wb") as handle:
        shutil.copyfileobj(file.file, handle)
    return target


async def extract_pages(
    path: Path,
    settings: Settings,
    ocr_service: OcrService,
    ocr_engine: OcrEngineName,
    language: str,
) -> tuple[str, str, str, list[dict]]:
    document_id = _file_hash(path)
    filename = path.name
    mimetype = mimetypes.guess_type(path.name)[0] or "application/octet-stream"
    suffix = path.suffix.lower()

    if suffix in TEXT_SUFFIXES:
        return document_id, filename, mimetype, [{"page": 1, "text": path.read_text(encoding="utf-8", errors="replace"), "ocr_engine": None}]

    if suffix == ".docx":
        text = _extract_docx_text(path)
        return document_id, filename, mimetype, [{"page": 1, "text": text, "ocr_engine": None}]

    if suffix == ".pdf":
        pages = []
        if ocr_engine == "none":
            pages = _extract_pdf_text(path)
        if not pages or ocr_engine != "none":
            pages = await _ocr_pdf(path, settings, ocr_service, ocr_engine, language)
        return document_id, filename, mimetype, pages

    if suffix in IMAGE_SUFFIXES:
        image_path = _normalize_image(path)
        text, engine_used = await ocr_service.extract_text(image_path, ocr_engine, language)
        return document_id, filename, mimetype, [{"page": 1, "text": text, "ocr_engine": engine_used}]

    text = path.read_text(encoding="utf-8", errors="replace")
    return document_id, filename, mimetype, [{"page": 1, "text": text, "ocr_engine": None}]


def _file_hash(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _extract_pdf_text(path: Path) -> list[dict]:
    reader = PdfReader(str(path))
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append({"page": index, "text": text, "ocr_engine": None})
    return pages


async def _ocr_pdf(
    path: Path,
    settings: Settings,
    ocr_service: OcrService,
    ocr_engine: OcrEngineName,
    language: str,
) -> list[dict]:
    image_dir = settings.upload_dir / f"{path.stem}_pages_{uuid.uuid4().hex}"
    image_dir.mkdir(parents=True, exist_ok=True)
    page_images = []

    with fitz.open(path) as pdf_document:
        for page_index in range(pdf_document.page_count):
            image_path = image_dir / f"page_{page_index + 1:04d}.png"
            image = _render_pdf_page(pdf_document, page_index, dpi=200)
            image.save(image_path, "PNG")
            page_images.append((page_index + 1, image_path))

    concurrency = max(1, settings.ocr_concurrency)
    semaphore = asyncio.Semaphore(concurrency)

    async def ocr_page(page_number: int, image_path: Path) -> dict:
        async with semaphore:
            text, engine_used = await ocr_service.extract_text(image_path, ocr_engine, language, markdown=True)
            return {"page": page_number, "text": text, "ocr_engine": engine_used}

    pages = await asyncio.gather(*(ocr_page(page_number, image_path) for page_number, image_path in page_images))
    return sorted(pages, key=lambda page: page["page"])


def _render_pdf_page(pdf_document, page_index: int, dpi: int) -> Image.Image:
    page = pdf_document.load_page(page_index)
    zoom = dpi / 72
    pixmap = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
    return Image.open(io.BytesIO(pixmap.tobytes("png"))).convert("RGB")


def _normalize_image(path: Path) -> Path:
    normalized = path.with_suffix(".png")
    with Image.open(path) as image:
        image.convert("RGB").save(normalized, "PNG")
    return normalized


def write_parsed_output(parsed_root: Path, document_id: str, pages: list[dict]) -> tuple[Path, str]:
    document_dir = parsed_root / document_id
    pages_dir = document_dir / "pages"
    pages_dir.mkdir(parents=True, exist_ok=True)

    page_markdowns = []
    for index, page in enumerate(pages, start=1):
        page_number = page.get("page") or index
        text = (page.get("text") or "").strip()
        page_markdown = f"<!-- page {page_number} -->\n\n{text}"
        (pages_dir / f"page_{page_number:04d}.md").write_text(page_markdown, encoding="utf-8")
        page_markdowns.append(page_markdown)

    combined = "\n\n---\n\n".join(page_markdowns).strip()
    combined_path = document_dir / "combined.md"
    manifest_path = document_dir / "manifest.json"
    combined_path.write_text(combined, encoding="utf-8")
    manifest_path.write_text(
        json.dumps(
            {"document_id": document_id, "pages": len(pages), "combined_path": str(combined_path)},
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    return combined_path, combined[:2000]


def _extract_docx_text(path: Path) -> str:
    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append("\t".join(cells))
    return "\n".join(parts)
